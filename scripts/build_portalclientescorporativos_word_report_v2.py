#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_portalclientescorporativos_word_report_v2.py

Informe técnico mejorado con infografías para portalclientescorporativos.
Mejoras:
  - Descripciones enriquecidas basadas en análisis del código
  - Sin "Código (preview)"
  - Tablas SAP/HANA separadas de tablas SQL, listadas en viñetas
  - Infografías embebidas (PIL): gráfico de barras + pastel
  - Cajas de resumen coloreadas por handler/categoría
  - Lenguaje técnico pero accesible
"""

import json
import re
import io
import math
from collections import defaultdict
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ─── Rutas ─────────────────────────────────────────────────────────────────

RAW_JSON = Path(
    r"c:\Users\edelacruz\AppData\Roaming\Code\User\workspaceStorage"
    r"\6649269cc0399a69278f4be3179fa9ae\GitHub.copilot-chat\chat-session-resources"
    r"\81ab01c4-8428-4d65-aa34-11fd22da9743"
    r"\call_VYPNJe70RIvM34EYdP0bcCk6__vscode-1774875362646\content.json"
)
OUT_DOCX = Path(
    r"d:\edwinspire\OtrosProyectos\libOpenFusionAPI\docs"
    r"\portalclientescorporativos_informe_tecnico_v2.docx"
)

# ─── Paleta de colores ──────────────────────────────────────────────────────

HANDLER_RGB = {
    "SQL":  (0x15, 0x65, 0xC0),   # azul profundo - SQL Server
    "HANA": (0xE6, 0x51, 0x00),   # naranja SAP
    "JS":   (0x2E, 0x7D, 0x32),   # verde - orquestador
    "SOAP": (0x6A, 0x1B, 0x9A),   # púrpura - servicio SOAP
    "TEXT": (0x54, 0x6E, 0x7A),   # gris azul - contenido estático
    "MCP":  (0x00, 0x83, 0x8F),   # teal - agente IA
}
HANDLER_HEX = {k: "%02X%02X%02X" % v for k, v in HANDLER_RGB.items()}

CATEGORY_HEX = [
    "1565C0", "2E7D32", "F57F17", "6A1B9A",
    "00838F", "C62828", "4E342E",
]

HANDLER_LABEL = {
    "SQL":  "SQL Server",
    "HANA": "SAP / HANA",
    "JS":   "JavaScript (Orquestador)",
    "SOAP": "SOAP Web Service",
    "TEXT": "Contenido Estático",
    "MCP":  "MCP – Herramienta de Agente IA",
}

METHOD_LABEL = {
    "GET":    "Consulta (solo lectura)",
    "POST":   "Creación / Procesamiento",
    "PUT":    "Actualización",
    "DELETE": "Eliminación",
    "PATCH":  "Modificación parcial",
}

CATEGORY_ORDER = [
    "Administracion usuarios externos",
    "Administración de Usuarios Farma",
    "Administracion Corporativo - Creacion convenio",
    "Administración Corporativo - Cambio Dia Corte",
    "Administracion Corporativo - Resumen Convenio",
    "Administración Corporativo – Colaboradores",
    "Administración Corporativo – Actualización de Cupos",
]

RISK_LABEL = {"bajo": "BAJO ▼", "medio": "MEDIO ◆", "alto": "ALTO ▲"}
RISK_HEX   = {"bajo": "43A047", "medio": "FB8C00", "alto": "E53935"}

# ─── Utilidades generales ───────────────────────────────────────────────────

def uniq(seq):
    seen, out = set(), []
    for x in seq:
        if x and x not in seen:
            seen.add(x); out.append(x)
    return out


def clean_resource_from_call(path):
    r = str(path or "")
    r = r.replace("/api/portalclientescorporativos", "")
    r = re.sub(r"/auto$", "", r)
    return r


def categorize(resource):
    x = str(resource or "").lower()
    if "/colaboradores/cupos" in x or "/cupos/carga" in x:
        return "Administración Corporativo – Actualización de Cupos"
    if "/cambiodiacorte" in x or "/cambio_dia_corte" in x or "/bbdd/cortes/" in x or "/corte/ultimo" in x:
        return "Administración Corporativo - Cambio Dia Corte"
    if "/corporativo/resumen" in x or "/bbdd/corporativo/cartera" in x or "/sap/fbl5n" in x or "/sap/bp" in x or "/bbdd/corporativos" in x:
        return "Administracion Corporativo - Resumen Convenio"
    if "/colaboradores/" in x or x == "/colaboradores" or "/clientes/carga/antiguo" in x or "/clientes/modifcacion/antiguo" in x or "/estados/carga/antiguo" in x:
        return "Administración Corporativo – Colaboradores"
    if "/corporativo/ejecutivos" in x or "/bbdd/credito_cobrazas/usuarios" in x or "/ejecutivos/carga/masiva" in x:
        return "Administración de Usuarios Farma"
    if ("/corporativo" in x and "/resumen" not in x and "/cambio_dia_corte" not in x and "/ejecutivos" not in x) or "/bbdd/corporativo" in x:
        return "Administracion Corporativo - Creacion convenio"
    if "/usuario" in x or "/usuarios/externos" in x or "/easyseguridad" in x or "/email/" in x:
        return "Administracion usuarios externos"
    return "No clasificado"


def extract_internal_calls(code):
    calls = re.findall(r"/api/portalclientescorporativos/[A-Za-z0-9_\-/]+", str(code or ""))
    return uniq([clean_resource_from_call(c) for c in calls])


def extract_cte_names(sql):
    s = str(sql or "")
    ctes = []
    for m in re.finditer(r"\bWITH\s+([A-Za-z_][A-Za-z0-9_]*)\s+AS\s*\(", s, re.IGNORECASE):
        ctes.append(m.group(1).lower())
    for m in re.finditer(r"\),\s*([A-Za-z_][A-Za-z0-9_]*)\s+AS\s*\(", s, re.IGNORECASE):
        ctes.append(m.group(1).lower())
    return set(ctes)


def extract_tables(sql):
    s = str(sql or "")
    cte_names = extract_cte_names(s)
    skip = {"SELECT", "VALUES", "SET", "BEGIN", "END", "UPDATE", "INTO", "NULL", "WITH"}
    tables = []
    pat = re.compile(r"\b(?:FROM|JOIN|UPDATE|INTO|DELETE\s+FROM|MERGE\s+INTO)\s+([\[\]A-Za-z0-9_$.#@]+)", re.IGNORECASE)
    for m in pat.finditer(s):
        t = m.group(1).strip().rstrip(".")
        if t.lower() in cte_names: continue
        if len(t) <= 1: continue
        if t.upper() in skip: continue
        if t.startswith("@") or t.startswith("#"): continue
        tables.append(t)
    return uniq(tables)


def extract_params(code):
    params = re.findall(r'\$([A-Za-z_][A-Za-z0-9_]*)', str(code or ""))
    return uniq(params)[:10]


def detect_operations(code, handler):
    s = str(code or "")
    ops = []
    if handler in ("SQL", "HANA"):
        if re.search(r'\bSELECT\b', s, re.IGNORECASE): ops.append("consulta de datos")
        if re.search(r'\bINSERT\b', s, re.IGNORECASE): ops.append("inserción de registros")
        if re.search(r'\bUPDATE\b', s, re.IGNORECASE): ops.append("actualización de registros")
        if re.search(r'\bDELETE\b', s, re.IGNORECASE): ops.append("eliminación de registros")
        if re.search(r'\bMERGE\b',  s, re.IGNORECASE): ops.append("sincronización (MERGE)")
        if re.search(r'\bTHROW\b|\bRAISERROR\b', s, re.IGNORECASE): ops.append("validación con errores controlados")
        if re.search(r'\bBEGIN\s+TRAN|\bTRANSACTION\b', s, re.IGNORECASE): ops.append("transacción atómica")
    elif handler == "JS":
        if re.search(r'xlsx|workbook|XLSX', s): ops.append("generación de archivo Excel")
        if re.search(r'fetch\s*\(', s): ops.append("coordinación de múltiples servicios internos")
        if re.search(r'email|sendMail|nodemailer', s, re.IGNORECASE): ops.append("envío de correo electrónico")
        if re.search(r'pdf|pdfkit|PDFDocument', s, re.IGNORECASE): ops.append("generación de PDF")
        if re.search(r'masiv|bulk|lote', s, re.IGNORECASE): ops.append("procesamiento masivo de registros")
        if re.search(r'FormData|multipart', s, re.IGNORECASE): ops.append("procesamiento de archivos")
    elif handler == "SOAP":
        op = re.search(r'<([A-Za-z][A-Za-z0-9]+)[\s">]', s)
        if op: ops.append(f"operación SOAP '{op.group(1)}'")
    return ops


def build_rich_description(ep, internal_calls, sql_tables, hana_tables):
    code    = str(ep.get("code") or "")
    handler = str(ep.get("handler") or "")
    resource = str(ep.get("resource") or "")
    method  = str(ep.get("method") or "")
    desc    = str(ep.get("description") or "").strip()

    parts = []

    # 1) Base description
    if desc:
        cleaned = " ".join(desc.replace("\r", " ").replace("\n", " ").split())
        parts.append(cleaned)
    else:
        r = resource.replace("/", " ").replace("_", " ").strip()
        if "plantilla" in resource.lower():
            parts.append(f"Genera una plantilla descargable para el proceso de {r}.")
        elif "masiva" in resource.lower() or "masivo" in resource.lower():
            parts.append(f"Procesa la carga masiva de registros para {r}.")
        elif "creacion" in resource.lower() or "creación" in resource.lower():
            parts.append(f"Registra un nuevo elemento en el proceso de {r}.")
        elif "bbdd" in resource.lower():
            parts.append(f"Ejecuta una operación de base de datos para {r}.")
        elif "sap" in resource.lower() or "fbl5n" in resource.lower() or "bp" in resource.lower():
            parts.append(f"Accede al sistema SAP para obtener información de {r}.")
        else:
            parts.append(f"Servicio técnico para el proceso de {r}.")

    # 2) Operations
    ops = detect_operations(code, handler)
    if ops:
        parts.append(f"Operaciones que realiza: {', '.join(ops)}.")

    # 3) Parameters
    params = extract_params(code)
    if params:
        parts.append(f"Parámetros de entrada detectados: {', '.join(params)}.")

    # 4) HTTP method intent
    m_map = {
        "GET":    "Al ser un método GET, solo consulta datos sin modificar el sistema.",
        "POST":   "Al ser un método POST, recibe datos del cliente y ejecuta el procesamiento.",
        "PUT":    "Al ser un método PUT, actualiza un recurso existente con los datos recibidos.",
        "DELETE": "Al ser un método DELETE, elimina o desactiva el recurso indicado.",
    }
    if method in m_map:
        parts.append(m_map[method])

    # 5) Handler-specific context
    if handler == "JS" and internal_calls:
        parts.append(
            f"Este endpoint actúa como orquestador: coordina {len(internal_calls)} "
            f"servicio(s) interno(s) del mismo dominio para completar la operación."
        )
    elif handler == "SOAP":
        parts.append(
            "Se conecta a un servicio web SOAP legado. "
            "La lógica de negocio reside en el servicio externo; este endpoint actúa como puente."
        )
    elif handler == "TEXT":
        parts.append(
            "Entrega un recurso de texto estático (plantilla de correo, HTML u otro documento fijo) "
            "sin lógica dinámica."
        )
    elif handler == "MCP":
        parts.append(
            "Expone una capacidad como herramienta para el agente de inteligencia artificial "
            "a través del protocolo MCP."
        )

    # 6) DB impact brief
    if sql_tables:
        parts.append(
            f"Impacta directa o indirectamente {len(sql_tables)} tabla(s)/vista(s) "
            f"en SQL Server: {', '.join(sql_tables[:4])}{'...' if len(sql_tables) > 4 else ''}."
        )
    if hana_tables:
        parts.append(
            f"Accede a {len(hana_tables)} objeto(s) en SAP/HANA: "
            f"{', '.join(hana_tables[:4])}{'...' if len(hana_tables) > 4 else ''}."
        )

    return " ".join(parts)

# ─── Carga y construcción del dataset ──────────────────────────────────────

def load_endpoints():
    payload = json.loads(RAW_JSON.read_text(encoding="utf-8"))
    app = payload[0] if isinstance(payload, list) else payload
    return app, app.get("endpoints", [])


def build_dataset(app, eps):
    sql_tables_by_res  = {}
    hana_tables_by_res = {}
    handler_by_res     = {}

    for e in eps:
        res = e.get("resource")
        h   = e.get("handler")
        handler_by_res[res] = h
        if h == "SQL":
            sql_tables_by_res[res]  = extract_tables(e.get("code"))
        elif h == "HANA":
            hana_tables_by_res[res] = extract_tables(e.get("code"))

    rows = []
    for e in eps:
        res     = e.get("resource")
        handler = e.get("handler")
        method  = e.get("method")

        internal_calls = extract_internal_calls(e.get("code")) if handler == "JS" else []

        # Direct tables
        direct_sql  = list(sql_tables_by_res.get(res,  [])) if handler == "SQL"  else []
        direct_hana = list(hana_tables_by_res.get(res, [])) if handler == "HANA" else []

        # Indirect tables (via JS calls)
        indirect_sql  = []
        indirect_hana = []
        for c in internal_calls:
            indirect_sql.extend(sql_tables_by_res.get(c, []))
            indirect_hana.extend(hana_tables_by_res.get(c, []))

        sql_t  = uniq(direct_sql  + indirect_sql)
        hana_t = uniq(direct_hana + indirect_hana)
        all_t  = uniq(sql_t + hana_t)

        row = {
            "environment":     e.get("environment"),
            "method":          method,
            "resource":        res,
            "handler":         handler,
            "category":        categorize(res),
            "internal_calls":  internal_calls,
            "sql_tables":      sql_t,
            "hana_tables":     hana_t,
            "all_tables":      all_t,
            "is_direct_sql":   bool(direct_sql),
            "is_direct_hana":  bool(direct_hana),
            "is_indirect":     bool(indirect_sql or indirect_hana),
        }
        # Build rich description
        row["description"] = build_rich_description(e, internal_calls, sql_t, hana_t)
        rows.append(row)

    grouped = defaultdict(list)
    for r in rows:
        grouped[r["category"]].append(r)
    for cat in grouped:
        grouped[cat].sort(key=lambda x: (x["resource"] or "", x["method"] or ""))

    # Traceability
    traceability = []
    for r in rows:
        if r["handler"] != "JS":
            continue
        if not r["internal_calls"]:
            traceability.append({
                "source":         f"{r['method']} {r['resource']}",
                "target":         "(sin llamadas internas detectadas)",
                "target_handler": "-",
                "sql_tables":     [],
                "hana_tables":    [],
            })
            continue
        for c in r["internal_calls"]:
            th = handler_by_res.get(c, "No encontrado")
            traceability.append({
                "source":         f"{r['method']} {r['resource']}",
                "target":         c,
                "target_handler": th,
                "sql_tables":     sql_tables_by_res.get(c, []),
                "hana_tables":    hana_tables_by_res.get(c, []),
            })

    return grouped, rows, traceability


def executive_summary_for_category(cat, rows):
    handlers    = defaultdict(int)
    total_calls = 0
    all_sql     = set()
    all_hana    = set()

    for r in rows:
        handlers[r["handler"]] += 1
        total_calls += len(r["internal_calls"])
        all_sql.update(r["sql_tables"])
        all_hana.update(r["hana_tables"])

    handler_text = " · ".join(
        f"{HANDLER_LABEL.get(k, k)}: {v}" for k, v in sorted(handlers.items())
    )

    risk = "bajo"
    if len(all_sql) + len(all_hana) > 5 or total_calls > 10:
        risk = "alto"
    elif len(all_sql) + len(all_hana) > 2 or total_calls > 4:
        risk = "medio"

    lines = [
        f"Cantidad de endpoints: {len(rows)}",
        f"Tecnologías involucradas: {handler_text}",
    ]
    if all_sql:
        lines.append(f"Tablas/vistas SQL Server ({len(all_sql)}): " + ", ".join(sorted(all_sql)))
    if all_hana:
        lines.append(f"Objetos SAP/HANA ({len(all_hana)}): " + ", ".join(sorted(all_hana)))
    if total_calls:
        lines.append(f"Llamadas internas detectadas: {total_calls}")
    lines.append(f"Nivel de acoplamiento estimado: {RISK_LABEL[risk]}")

    return lines, risk

# ─── Infografías PIL ────────────────────────────────────────────────────────

def _try_font(path, size):
    try:
        return ImageFont.truetype(path, size)
    except Exception:
        return None

def get_font(size=14, bold=False):
    if not HAS_PIL:
        return None
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf"  if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    for c in candidates:
        f = _try_font(c, size)
        if f:
            return f
    return ImageFont.load_default()


def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def pil_horizontal_bar_chart(labels, values, colors_hex, title, subtitle=""):
    """Generate a horizontal bar chart as PNG bytes."""
    if not HAS_PIL:
        return None

    bar_h   = 40
    padding = 60
    label_w = 320
    bar_area = 480
    width   = label_w + bar_area + padding * 2
    height  = padding + 50 + len(labels) * (bar_h + 14) + padding

    img  = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    f_title  = get_font(18, bold=True)
    f_sub    = get_font(13)
    f_label  = get_font(14)
    f_val    = get_font(13, bold=True)

    # Title
    draw.text((padding, padding // 2), title, fill="#212121", font=f_title)
    if subtitle:
        draw.text((padding, padding // 2 + 26), subtitle, fill="#757575", font=f_sub)

    max_val = max(values) if values else 1
    y = padding + 50

    for i, (label, val, col_hex) in enumerate(zip(labels, values, colors_hex)):
        bar_len = int((val / max_val) * (bar_area - 30))
        rgb     = hex_to_rgb(col_hex)

        # Label (right-aligned inside label area)
        bbox = draw.textbbox((0, 0), label, font=f_label)
        text_w = bbox[2] - bbox[0]
        tx = padding + label_w - text_w - 8
        draw.text((tx, y + (bar_h - 16) // 2), label, fill="#212121", font=f_label)

        # Bar background
        draw.rounded_rectangle(
            [padding + label_w, y, padding + label_w + bar_area - 30, y + bar_h],
            radius=4, fill="#F5F5F5"
        )
        # Bar fill
        if bar_len > 0:
            draw.rounded_rectangle(
                [padding + label_w, y, padding + label_w + bar_len, y + bar_h],
                radius=4, fill=rgb
            )

        # Value label
        draw.text(
            (padding + label_w + bar_len + 6, y + (bar_h - 14) // 2),
            str(val), fill="#212121", font=f_val
        )
        y += bar_h + 14

    # Thin bottom line
    draw.line([(padding, height - 20), (width - padding, height - 20)], fill="#E0E0E0", width=1)

    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    return buf.getvalue()


def pil_pie_chart(labels, values, colors_hex, title):
    """Generate a pie chart as PNG bytes."""
    if not HAS_PIL:
        return None

    size     = 520
    cx, cy   = size // 2, size // 2 - 20
    radius   = 160
    img      = Image.new("RGB", (size, size + 80), "white")
    draw     = ImageDraw.Draw(img)

    f_title  = get_font(17, bold=True)
    f_legend = get_font(13)
    f_pct    = get_font(12, bold=True)

    draw.text((20, 14), title, fill="#212121", font=f_title)

    total = sum(values) or 1
    start = -90

    slices = []
    for val, col_hex in zip(values, colors_hex):
        sweep = 360 * val / total
        slices.append((start, sweep, col_hex))
        start += sweep

    for (s, sw, ch) in slices:
        rgb = hex_to_rgb(ch)
        draw.pieslice(
            [cx - radius, cy - radius, cx + radius, cy + radius],
            start=s, end=s + sw,
            fill=rgb, outline="white",
        )
    # Inner white circle for donut look
    inner_r = radius * 0.45
    draw.ellipse(
        [cx - inner_r, cy - inner_r, cx + inner_r, cy + inner_r],
        fill="white"
    )

    # Legend below
    lx, ly = 20, size - 50
    cols_per_row = 2
    for idx, (label, val, col_hex) in enumerate(zip(labels, values, colors_hex)):
        rgb  = hex_to_rgb(col_hex)
        pct  = f"{val / total * 100:.0f}%"
        col  = idx % cols_per_row
        row  = idx // cols_per_row
        x    = lx + col * (size // 2)
        y    = ly + row * 22
        draw.rectangle([x, y + 3, x + 14, y + 14], fill=rgb)
        draw.text((x + 18, y), f"{label} ({val} – {pct})", fill="#424242", font=f_legend)

    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    return buf.getvalue()


def pil_category_stats_bar(handler_counts):
    """Small horizontal stacked info bar for a category (handler breakdown)."""
    if not HAS_PIL:
        return None

    width, height = 720, 54
    img  = Image.new("RGB", (width, height), "#FAFAFA")
    draw = ImageDraw.Draw(img)
    f    = get_font(12, bold=True)

    total  = sum(handler_counts.values()) or 1
    bar_h  = 28
    bar_y  = 14
    bar_x  = 10
    bar_w  = width - 20

    # Draw stacked bar
    x_cursor = bar_x
    segments  = []
    for handler, cnt in sorted(handler_counts.items()):
        seg_w = int((cnt / total) * bar_w)
        if seg_w > 0:
            rgb = HANDLER_RGB.get(handler, (0x90, 0x90, 0x90))
            draw.rectangle([x_cursor, bar_y, x_cursor + seg_w, bar_y + bar_h], fill=rgb)
            segments.append((x_cursor, seg_w, handler, cnt, rgb))
            x_cursor += seg_w

    # Labels inside segments
    for (bx, bw, handler, cnt, rgb) in segments:
        label = f"{handler} ({cnt})"
        bbox = draw.textbbox((0, 0), label, font=f)
        tw = bbox[2] - bbox[0]
        if bw > tw + 8:
            tx = bx + (bw - tw) // 2
            # Contrast: decide white or dark text
            brightness = 0.299 * rgb[0] + 0.587 * rgb[1] + 0.114 * rgb[2]
            fc = "white" if brightness < 140 else "#212121"
            draw.text((tx, bar_y + 7), label, fill=fc, font=f)

    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    return buf.getvalue()

# ─── Helpers de estilo Word ─────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    color = hex_color.lstrip("#")
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color.upper())
    # Remove previous shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="BDBDBD", sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def p_run(paragraph, text, bold=False, italic=False,
          size_pt=None, color_hex=None, font_name=None):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        h = color_hex.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    if font_name:
        run.font.name = font_name
    return run


def add_bullet(doc, text, level=0, color_hex=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color_hex:
        h = color_hex.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), "E0E0E0")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_handler_badge_row(doc, method, resource, handler):
    """Add a coloured header row for each endpoint."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    h_hex = HANDLER_HEX.get(handler, "9E9EDD")

    # Left cell: method + handler
    c0 = tbl.rows[0].cells[0]
    set_cell_bg(c0, h_hex)
    c0.width = Cm(4.5)
    cp = c0.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run(cp, f"{method}  ", bold=True, size_pt=11, color_hex="FFFFFF")
    p_run(cp, f"[ {HANDLER_LABEL.get(handler, handler)} ]", bold=False, size_pt=9, color_hex="FFFFFF")

    # Right cell: resource path
    c1 = tbl.rows[0].cells[1]
    set_cell_bg(c1, "F5F5F5")
    rp = c1.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_run(rp, resource, bold=True, size_pt=10.5, color_hex="212121", font_name="Consolas")

    # Remove outer borders
    for cell in (c0, c1):
        set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="BDBDBD", sz=6)
    return tbl


def add_info_box(doc, label_color_hex, icon_label, body_lines):
    """Coloured info box: icon on left, content lines on right."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"

    c0 = tbl.rows[0].cells[0]
    set_cell_bg(c0, label_color_hex)
    c0.width = Cm(1.2)
    cp = c0.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run(cp, icon_label, bold=True, size_pt=9, color_hex="FFFFFF")

    c1 = tbl.rows[0].cells[1]
    set_cell_bg(c1, "FAFAFA")
    first = True
    for line in body_lines:
        p = c1.paragraphs[0] if first else c1.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.add_run(line).font.size = Pt(10)

    for cell in (c0, c1):
        set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="E0E0E0", sz=4)
    return tbl


def add_tables_section(doc, sql_tables, hana_tables, is_direct_sql, is_direct_hana, is_indirect):
    """Render SQL and HANA tables as separate visual blocks."""
    if not sql_tables and not hana_tables:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p_run(p, "Sin impacto visible en base de datos detectado desde el código.", italic=True, color_hex="9E9E9E", size_pt=9.5)
        return

    # SQL block
    if sql_tables:
        origin_note = ""
        if is_direct_sql and is_indirect:
            origin_note = " (directas + trazabilidad JS)"
        elif is_direct_sql:
            origin_note = " (directas)"
        elif is_indirect:
            origin_note = " (por trazabilidad desde JS)"

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells[0]
        set_cell_bg(hdr, HANDLER_HEX["SQL"])
        hdr.width = Cm(1.6)
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_run(hp, "SQL\nServer", bold=True, size_pt=9, color_hex="FFFFFF")

        body = tbl.rows[0].cells[1]
        set_cell_bg(body, "EEF4FB")
        bp = body.paragraphs[0]
        p_run(bp, f"Tablas / Vistas SQL Server{origin_note}", bold=True, size_pt=9.5, color_hex="1565C0")
        for t in sql_tables:
            nb = body.add_paragraph()
            nb.paragraph_format.left_indent = Cm(0.3)
            nb.paragraph_format.space_before = Pt(0)
            nb.paragraph_format.space_after  = Pt(0)
            p_run(nb, "▸ ", bold=True, size_pt=9, color_hex="1565C0")
            p_run(nb, t, size_pt=9.5, font_name="Consolas")

        for cell in (hdr, body):
            set_cell_border(cell, color="BBDEFB", sz=4)

    # HANA block
    if hana_tables:
        origin_note = ""
        if is_direct_hana and is_indirect:
            origin_note = " (directas + trazabilidad JS)"
        elif is_direct_hana:
            origin_note = " (directas)"
        elif is_indirect:
            origin_note = " (por trazabilidad desde JS)"

        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = "Table Grid"
        hdr2 = tbl2.rows[0].cells[0]
        set_cell_bg(hdr2, HANDLER_HEX["HANA"])
        hdr2.width = Cm(1.6)
        hp2 = hdr2.paragraphs[0]
        hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_run(hp2, "SAP\nHANA", bold=True, size_pt=9, color_hex="FFFFFF")

        body2 = tbl2.rows[0].cells[1]
        set_cell_bg(body2, "FEF3E2")
        bp2 = body2.paragraphs[0]
        p_run(bp2, f"Objetos SAP / HANA{origin_note}", bold=True, size_pt=9.5, color_hex="E65100")
        for t in hana_tables:
            nb = body2.add_paragraph()
            nb.paragraph_format.left_indent = Cm(0.3)
            nb.paragraph_format.space_before = Pt(0)
            nb.paragraph_format.space_after  = Pt(0)
            p_run(nb, "▸ ", bold=True, size_pt=9, color_hex="E65100")
            p_run(nb, t, size_pt=9.5, font_name="Consolas")

        for cell2 in (hdr2, body2):
            set_cell_border(cell2, color="FFCC80", sz=4)


def add_internal_calls_block(doc, calls, handler_by_res):
    if not calls:
        return
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells[0]
    set_cell_bg(hdr, HANDLER_HEX["JS"])
    hdr.width = Cm(1.6)
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run(hp, "Llama­das\nInter­nas", bold=True, size_pt=9, color_hex="FFFFFF")

    body = tbl.rows[0].cells[1]
    set_cell_bg(body, "F1F8F1")
    first = True
    for c in calls:
        th = handler_by_res.get(c, "?")
        badge = HANDLER_HEX.get(th, "9E9E9E")
        p = body.paragraphs[0] if first else body.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p_run(p, "→ ", bold=True, size_pt=9.5, color_hex="2E7D32")
        p_run(p, c, size_pt=9.5, font_name="Consolas")
        p_run(p, f"  [{HANDLER_LABEL.get(th, th)}]", italic=True, size_pt=8.5, color_hex="616161")

    for cell in (hdr, body):
        set_cell_border(cell, color="C8E6C9", sz=4)


def add_executive_summary_box(doc, lines, risk):
    """Coloured box for the executive summary of a category."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells[0]
    set_cell_bg(hdr, RISK_HEX[risk])
    hdr.width = Cm(1.6)
    hp  = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run(hp, "RESU­MEN\nEJEC­U­TIVO", bold=True, size_pt=8.5, color_hex="FFFFFF")

    body = tbl.rows[0].cells[1]
    set_cell_bg(body, "F9F9F9")
    first = True
    for line in lines:
        p = body.paragraphs[0] if first else body.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.add_run(line).font.size = Pt(10)

    for cell in (hdr, body):
        set_cell_border(cell, color="BDBDBD", sz=6)

# ─── Escritura del documento ────────────────────────────────────────────────

def write_doc(app, grouped, rows, traceability):
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Global font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────────────────
    t = doc.add_heading("Informe Técnico de Endpoints", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t2 = doc.add_heading("Portal Clientes Corporativos", 1)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(6)
    p_run(meta, f"Fecha de generación: {date.today().strftime('%d de %B de %Y')}   |   "
                f"Total endpoints analizados: {len(rows)}   |   "
                f"Fuente: MCP openfusionapi_system_local", size_pt=9.5, color_hex="616161")

    doc.add_page_break()

    # ── RESUMEN EJECUTIVO GLOBAL con infografías ──────────────────────────
    doc.add_heading("Resumen Ejecutivo Global", level=1)

    # Stats text
    handler_count = defaultdict(int)
    for r in rows:
        handler_count[r["handler"]] += 1

    by_cat  = {k: len(grouped.get(k, [])) for k in CATEGORY_ORDER}
    db_eps  = sum(1 for r in rows if r["all_tables"])
    all_sql = set()
    all_hana = set()
    for r in rows:
        all_sql.update(r["sql_tables"])
        all_hana.update(r["hana_tables"])

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    intro.add_run(
        f"Este documento analiza los {len(rows)} endpoints del aplicativo "
        f"portalclientescorporativos, clasificados en {len(CATEGORY_ORDER)} categorías funcionales. "
        f"Se identificaron {db_eps} endpoint(s) con impacto en base de datos: "
        f"{len(all_sql)} tabla(s)/vista(s) en SQL Server y {len(all_hana)} objeto(s) en SAP/HANA."
    )

    # Bar chart: endpoints per category
    if HAS_PIL:
        cat_labels = [c.split("–")[-1].split("-")[-1].strip() for c in CATEGORY_ORDER]
        cat_vals   = [by_cat.get(c, 0) for c in CATEGORY_ORDER]
        bar_png    = pil_horizontal_bar_chart(
            cat_labels, cat_vals, CATEGORY_HEX,
            title="Endpoints por categoría funcional",
            subtitle=f"Total: {len(rows)} endpoints analizados"
        )
        if bar_png:
            doc.add_paragraph().paragraph_format.space_before = Pt(4)
            doc.add_picture(io.BytesIO(bar_png), width=Inches(6.2))
            cp = doc.paragraphs[-1]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Pie chart: handler distribution
    if HAS_PIL:
        h_labels = list(handler_count.keys())
        h_vals   = [handler_count[k] for k in h_labels]
        h_colors = [HANDLER_HEX.get(k, "9E9E9E") for k in h_labels]
        pie_png  = pil_pie_chart(
            [HANDLER_LABEL.get(k, k) for k in h_labels],
            h_vals, h_colors,
            title="Distribución por tipo de handler"
        )
        if pie_png:
            doc.add_picture(io.BytesIO(pie_png), width=Inches(5.2))
            cp2 = doc.paragraphs[-1]
            cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── CATEGORÍAS ────────────────────────────────────────────────────────
    handler_by_res_local = {}
    for r in rows:
        handler_by_res_local[r["resource"]] = r["handler"]

    for cat_idx, cat in enumerate(CATEGORY_ORDER):
        cat_rows = grouped.get(cat, [])
        cat_hex  = CATEGORY_HEX[cat_idx % len(CATEGORY_HEX)]

        # Category heading
        h = doc.add_heading("", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_run(h, f"{cat}  ", bold=True)
        p_run(h, f"({len(cat_rows)} endpoints)", bold=False, size_pt=11, color_hex="757575")

        # Category stacked bar infographic
        if HAS_PIL and cat_rows:
            h_counts = defaultdict(int)
            for r in cat_rows:
                h_counts[r["handler"]] += 1
            stats_png = pil_category_stats_bar(dict(h_counts))
            if stats_png:
                doc.add_picture(io.BytesIO(stats_png), width=Inches(6.2))
                pgr = doc.paragraphs[-1]
                pgr.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Each endpoint
        for r in cat_rows:
            # Header badge
            add_handler_badge_row(doc, r["method"], r["resource"], r["handler"])

            # Description
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent  = Cm(0.4)
            dp.paragraph_format.space_before = Pt(2)
            dp.paragraph_format.space_after  = Pt(4)
            dp.add_run(r["description"]).font.size = Pt(10)

            # Tables: SQL and/or HANA
            add_tables_section(
                doc,
                r["sql_tables"], r["hana_tables"],
                r["is_direct_sql"], r["is_direct_hana"], r["is_indirect"]
            )

            # Internal calls (JS only)
            if r["internal_calls"]:
                add_internal_calls_block(doc, r["internal_calls"], handler_by_res_local)

            # Small spacer
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after  = Pt(6)

        # Category executive summary
        add_horizontal_rule(doc)
        sum_lines, risk = executive_summary_for_category(cat, cat_rows)
        add_executive_summary_box(doc, sum_lines, risk)
        doc.add_page_break()

    # ── ENDPOINTS NO CLASIFICADOS ─────────────────────────────────────────
    no_class = sorted(grouped.get("No clasificado", []), key=lambda x: (x["resource"] or ""))
    if no_class:
        doc.add_heading(f"Anexo A — Endpoints no clasificados ({len(no_class)})", level=1)
        for r in no_class:
            add_handler_badge_row(doc, r["method"], r["resource"], r["handler"])
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Cm(0.4)
            dp.add_run(r["description"]).font.size = Pt(10)
            add_tables_section(doc, r["sql_tables"], r["hana_tables"],
                               r["is_direct_sql"], r["is_direct_hana"], r["is_indirect"])
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        doc.add_page_break()

    # ── INFORME DE TRAZABILIDAD ──────────────────────────────────────────
    doc.add_heading("Informe de Trazabilidad — Endpoints JS", level=1)
    doc.add_paragraph(
        "Matriz de trazabilidad de los endpoints JavaScript (orquestadores) "
        "hacia sus dependencias internas. Se indica el handler del endpoint destino "
        "y las tablas/objetos de base de datos que impacta, diferenciando SQL Server de SAP/HANA."
    )

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"

    # Headers
    hdrs = ["Origen (JS)", "Ruta destino", "Handler destino", "Tablas SQL Server", "Objetos SAP/HANA"]
    for i, hdr_text in enumerate(hdrs):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "37474F")
        set_cell_border(c, color="616161", sz=4)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_run(p, hdr_text, bold=True, size_pt=9, color_hex="FFFFFF")

    for tr_idx, tr in enumerate(traceability):
        row_cells = tbl.add_row().cells
        row_cells[0].text = tr["source"]
        row_cells[1].text = tr["target"]
        row_cells[2].text = HANDLER_LABEL.get(tr["target_handler"], str(tr["target_handler"]))

        # SQL tables column
        sql_t = tr.get("sql_tables", [])
        if sql_t:
            p4 = row_cells[3].paragraphs[0]
            for i, t in enumerate(sql_t):
                if i > 0:
                    p4 = row_cells[3].add_paragraph()
                p4.paragraph_format.space_before = Pt(0)
                p4.paragraph_format.space_after  = Pt(0)
                p_run(p4, t, size_pt=9, font_name="Consolas")
        else:
            row_cells[3].paragraphs[0].add_run("—").font.size = Pt(9)

        # HANA tables column
        hana_t = tr.get("hana_tables", [])
        if hana_t:
            p5 = row_cells[4].paragraphs[0]
            for i, t in enumerate(hana_t):
                if i > 0:
                    p5 = row_cells[4].add_paragraph()
                p5.paragraph_format.space_before = Pt(0)
                p5.paragraph_format.space_after  = Pt(0)
                p_run(p5, t, size_pt=9, font_name="Consolas")
        else:
            row_cells[4].paragraphs[0].add_run("—").font.size = Pt(9)

        # Style rows with alternating bg
        row_bg = "F9FAFB" if (tr_idx % 2 == 0) else "FFFFFF"
        for cell in row_cells:
            set_cell_bg(cell, row_bg)
            set_cell_border(cell, color="E0E0E0", sz=4)


    doc.add_page_break()

    # ── GLOSARIO DE HANDLERS ──────────────────────────────────────────────
    doc.add_heading("Anexo B — Glosario de Tipos de Handler", level=1)
    glosario = [
        ("SQL Server", "SQL",
         "Endpoint que ejecuta una consulta o instrucción directamente en la base de datos "
         "SQL Server del sistema. Puede realizar selecciones, inserciones, actualizaciones o "
         "eliminaciones de registros."),
        ("SAP / HANA", "HANA",
         "Endpoint que ejecuta una consulta sobre la base de datos SAP HANA. "
         "Las tablas son objetos del ecosistema SAP (ej. SAPABAP1.UKM_ITEM). "
         "Se usa para obtener datos financieros y de crédito desde SAP."),
        ("JavaScript (Orquestador)", "JS",
         "Endpoint que implementa lógica de negocio en JavaScript del lado del servidor. "
         "Típicamente coordina múltiples llamadas internas, procesa archivos Excel, "
         "genera plantillas o valida datos antes de persistirlos."),
        ("SOAP Web Service", "SOAP",
         "Endpoint que actúa como puente hacia un servicio web SOAP legado. "
         "La lógica de negocio reside en el servicio externo; este endpoint sólo adapta "
         "el contrato de la petición/respuesta."),
        ("Contenido Estático", "TEXT",
         "Endpoint que retorna un recurso de texto fijo: una plantilla de correo HTML, "
         "un fragmento XML o cualquier contenido que no requiere procesamiento dinámico."),
        ("MCP – Herramienta IA", "MCP",
         "Endpoint que expone una herramienta para ser consumida por un agente de "
         "inteligencia artificial a través del protocolo Model Context Protocol (MCP)."),
    ]

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    for i, hdr_text in enumerate(["Tipo de Handler", "Código", "Descripción"]):
        c = tbl2.rows[0].cells[i]
        set_cell_bg(c, "37474F")
        p = c.paragraphs[0]
        p_run(p, hdr_text, bold=True, size_pt=9.5, color_hex="FFFFFF")

    for label, code, description in glosario:
        r = tbl2.add_row().cells
        r[0].text = label
        r[1].text = code
        p_run(r[1].paragraphs[0], "", bold=False)
        set_cell_bg(r[1], HANDLER_HEX.get(code, "9E9E9E"))
        p3 = r[1].paragraphs[0]
        p3.clear()
        p_run(p3, code, bold=True, size_pt=9, color_hex="FFFFFF")
        r[2].text = description
        for cell in r:
            set_cell_border(cell, color="E0E0E0", sz=4)

    # ── GUARDAR ───────────────────────────────────────────────────────────
    doc.save(str(OUT_DOCX))
    print(f"Documento guardado: {OUT_DOCX}")
    if HAS_PIL:
        print("Infografías PIL: INCLUIDAS")
    else:
        print("ADVERTENCIA: Pillow no disponible, sin infografías de imagen")


if __name__ == "__main__":
    print("Cargando datos de endpoints...")
    app, eps = load_endpoints()
    print(f"  {len(eps)} endpoints cargados")
    print("Construyendo dataset...")
    grouped, rows, traceability = build_dataset(app, eps)
    print(f"  {len(rows)} filas, {len(traceability)} entradas de trazabilidad")
    print("Generando documento Word...")
    write_doc(app, grouped, rows, traceability)
