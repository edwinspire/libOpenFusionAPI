import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt


RAW_JSON = Path(
    r"c:\Users\edelacruz\AppData\Roaming\Code\User\workspaceStorage\6649269cc0399a69278f4be3179fa9ae\GitHub.copilot-chat\chat-session-resources\81ab01c4-8428-4d65-aa34-11fd22da9743\call_VYPNJe70RIvM34EYdP0bcCk6__vscode-1774875362646\content.json"
)
OUT_DOCX = Path(r"d:\edwinspire\OtrosProyectos\libOpenFusionAPI\docs\portalclientescorporativos_informe_tecnico.docx")

CATEGORY_ORDER = [
    "Administracion usuarios externos",
    "Administración de Usuarios Farma",
    "Administracion Corporativo - Creacion convenio",
    "Administración Corporativo - Cambio Dia Corte",
    "Administracion Corporativo - Resumen Convenio",
    "Administración Corporativo – Colaboradores",
    "Administración Corporativo – Actualización de Cupos",
]


def uniq(seq):
    seen = set()
    out = []
    for x in seq:
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def clean_resource_from_call(path):
    r = str(path or "")
    r = r.replace("/api/portalclientescorporativos", "")
    r = re.sub(r"/auto$", "", r)
    return r


def categorize(resource):
    x = str(resource or "").lower()
    if "/colaboradores/cupos" in x or "/cupos/carga" in x:
        return "Administración Corporativo – Actualización de Cupos"
    if "/cambiodiacorte" in x or "/cambio_dia_corte" in x or "/bbdd/cortes/" in x or "/corte/ultimo" in x:
        return "Administración Corporativo - Cambio Dia Corte"
    if "/corporativo/resumen" in x or "/bbdd/corporativo/cartera" in x or "/sap/fbl5n" in x or "/sap/bp" in x or "/bbdd/corporativos" in x:
        return "Administracion Corporativo - Resumen Convenio"
    if "/colaboradores/" in x or x == "/colaboradores" or "/clientes/carga/antiguo" in x or "/clientes/modifcacion/antiguo" in x or "/estados/carga/antiguo" in x:
        return "Administración Corporativo – Colaboradores"
    if "/corporativo/ejecutivos" in x or "/bbdd/credito_cobrazas/usuarios" in x or "/ejecutivos/carga/masiva" in x:
        return "Administración de Usuarios Farma"
    if ("/corporativo" in x and "/resumen" not in x and "/cambio_dia_corte" not in x and "/ejecutivos" not in x) or "/bbdd/corporativo" in x:
        return "Administracion Corporativo - Creacion convenio"
    if "/usuario" in x or "/usuarios/externos" in x or "/easyseguridad" in x or "/email/" in x:
        return "Administracion usuarios externos"
    return "No clasificado"


def extract_internal_calls(code):
    s = str(code or "")
    calls = re.findall(r"/api/portalclientescorporativos/[A-Za-z0-9_\-/]+", s)
    return uniq([clean_resource_from_call(c) for c in calls])


def extract_cte_names(sql):
    s = str(sql or "")
    ctes = []
    for m in re.finditer(r"\bWITH\s+([A-Za-z_][A-Za-z0-9_]*)\s+AS\s*\(", s, flags=re.IGNORECASE):
        ctes.append(m.group(1).lower())
    for m in re.finditer(r"\),\s*([A-Za-z_][A-Za-z0-9_]*)\s+AS\s*\(", s, flags=re.IGNORECASE):
        ctes.append(m.group(1).lower())
    return set(ctes)


def extract_tables(sql):
    s = str(sql or "")
    cte_names = extract_cte_names(s)
    tables = []
    pat = re.compile(r"\b(?:FROM|JOIN|UPDATE|INTO|DELETE\s+FROM|MERGE\s+INTO)\s+([\[\]A-Za-z0-9_$.]+)", flags=re.IGNORECASE)
    for m in pat.finditer(s):
        t = m.group(1).strip().rstrip(".")
        t_low = t.lower()
        if t_low in cte_names:
            continue
        if len(t) == 1:
            continue
        if t.upper() in {"SELECT", "VALUES", "SET", "BEGIN", "END", "UPDATE", "INTO"}:
            continue
        if t.startswith("@"):
            continue
        tables.append(t)
    return uniq(tables)


def infer_function_summary(ep):
    desc = str(ep.get("description") or "").strip()
    if desc:
        return desc.replace("\n", " ").strip()

    r = str(ep.get("resource") or "").lower()
    h = str(ep.get("handler") or "")

    if "/plantilla" in r:
        return "Genera una plantilla (generalmente HTML o Excel) para procesos masivos o comunicaciones."
    if "/masiva" in r:
        return "Procesa carga masiva desde archivo y orquesta validaciones/llamadas internas."
    if "/bbdd/" in r and h in {"SQL", "HANA"}:
        return "Ejecuta operación directa en base de datos para consulta o actualización."
    if h == "SOAP":
        return "Invoca un servicio SOAP legado para ejecutar la operación de negocio."
    if h == "JS":
        return "Orquesta lógica de negocio y consume endpoints internos del mismo dominio."
    if h == "TEXT":
        return "Entrega contenido estático (normalmente plantilla de correo)."
    if h == "MCP":
        return "Endpoint de soporte MCP para exponer herramientas al agente."
    return "Servicio técnico de soporte a procesos del dominio."


def executive_summary_for_category(cat, rows):
    handlers = defaultdict(int)
    total_calls = 0
    db_rows = 0
    db_tables = set()
    for r in rows:
        handlers[r["handler"]] += 1
        total_calls += len(r["internal_calls"])
        if r["all_tables"]:
            db_rows += 1
            for t in r["all_tables"]:
                db_tables.add(t)

    handler_text = ", ".join(f"{k}:{v}" for k, v in sorted(handlers.items()))
    return (
        f"En la categoría '{cat}' se identificaron {len(rows)} endpoints. "
        f"Distribución de handlers: {handler_text}. "
        f"{db_rows} endpoint(s) presentan interacción directa o indirecta con base de datos, "
        f"con {len(db_tables)} tabla(s)/vista(s) detectadas en total. "
        f"Se detectaron {total_calls} llamada(s) internas entre endpoints del dominio."
    )


def load_endpoints():
    payload = json.loads(RAW_JSON.read_text(encoding="utf-8"))
    app = payload[0] if isinstance(payload, list) else payload
    eps = app.get("endpoints", [])
    return app, eps


def build_dataset(app, eps):
    by_resource = {e.get("resource"): e for e in eps}

    direct_tables_by_resource = {}
    for e in eps:
        if e.get("handler") in {"SQL", "HANA"}:
            direct_tables_by_resource[e.get("resource")] = extract_tables(e.get("code"))

    rows = []
    for e in eps:
        resource = e.get("resource")
        handler = e.get("handler")
        internal_calls = extract_internal_calls(e.get("code")) if handler == "JS" else []

        direct_tables = extract_tables(e.get("code")) if handler in {"SQL", "HANA"} else []
        indirect_tables = []
        for c in internal_calls:
            indirect_tables.extend(direct_tables_by_resource.get(c, []))

        row = {
            "environment": e.get("environment"),
            "method": e.get("method"),
            "resource": resource,
            "handler": handler,
            "category": categorize(resource),
            "function_summary": infer_function_summary(e),
            "internal_calls": internal_calls,
            "direct_tables": uniq(direct_tables),
            "indirect_tables": uniq(indirect_tables),
            "all_tables": uniq(direct_tables + indirect_tables),
            "code_preview": (str(e.get("code") or "").strip().replace("\n", " "))[:160],
        }
        rows.append(row)

    grouped = defaultdict(list)
    for r in rows:
        grouped[r["category"]].append(r)

    traceability = []
    for r in rows:
        if r["handler"] != "JS":
            continue
        if not r["internal_calls"]:
            traceability.append({
                "source": f"{r['method']} {r['resource']}",
                "target": "(sin llamadas internas)",
                "target_handler": "-",
                "db_tables": "-",
            })
            continue

        for c in r["internal_calls"]:
            target = by_resource.get(c)
            if target:
                th = target.get("handler")
                tt = direct_tables_by_resource.get(c, [])
                traceability.append({
                    "source": f"{r['method']} {r['resource']}",
                    "target": c,
                    "target_handler": th,
                    "db_tables": ", ".join(tt) if tt else "No visible (JS/SOAP/TEXT/MCP o lógica interna)",
                })
            else:
                traceability.append({
                    "source": f"{r['method']} {r['resource']}",
                    "target": c,
                    "target_handler": "No encontrado",
                    "db_tables": "No trazable en inventario del app",
                })

    return grouped, rows, traceability


def write_doc(app, grouped, rows, traceability):
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("Informe Técnico de Endpoints - portalclientescorporativos", 0)
    p = doc.add_paragraph()
    p.add_run(f"Aplicación: {app.get('app')}\n")
    p.add_run(f"idapp: {app.get('idapp')}\n")
    p.add_run(f"Total endpoints analizados: {len(rows)}\n")
    p.add_run("Fuente: Servidor MCP openfusionapi_system_local (inventario + código endpoint).")

    for cat in CATEGORY_ORDER:
        cat_rows = sorted(grouped.get(cat, []), key=lambda x: (x["resource"] or "", x["method"] or ""))
        doc.add_heading(f"{cat} ({len(cat_rows)})", level=1)

        for r in cat_rows:
            doc.add_heading(f"{r['method']} {r['resource']}", level=2)
            doc.add_paragraph(f"Handler: {r['handler']}")
            doc.add_paragraph(f"Funcionalidad técnica: {r['function_summary']}")
            if r["internal_calls"]:
                doc.add_paragraph("Llamadas internas: " + ", ".join(r["internal_calls"]))
            else:
                doc.add_paragraph("Llamadas internas: No aplica")

            if r["all_tables"]:
                if r["direct_tables"] and r["indirect_tables"]:
                    origin = "(directas + indirectas por trazabilidad)"
                elif r["direct_tables"]:
                    origin = "(directas)"
                else:
                    origin = "(indirectas por trazabilidad)"
                doc.add_paragraph("Tablas/Vistas impactadas " + origin + ": " + ", ".join(r["all_tables"]))
            else:
                doc.add_paragraph("Tablas/Vistas impactadas: No visibles en el código del endpoint")

            if r["code_preview"]:
                doc.add_paragraph("Código (preview): " + r["code_preview"])

        doc.add_heading("Resumen Ejecutivo de la categoría", level=2)
        doc.add_paragraph(executive_summary_for_category(cat, cat_rows))

    no_class = sorted(grouped.get("No clasificado", []), key=lambda x: (x["resource"] or "", x["method"] or ""))
    if no_class:
        doc.add_heading(f"Anexo - Endpoints no clasificados ({len(no_class)})", level=1)
        for r in no_class:
            doc.add_paragraph(f"- {r['method']} {r['resource']} [{r['handler']}] - {r['function_summary']}")

    doc.add_heading("Informe de Trazabilidad", level=1)
    doc.add_paragraph(
        "Matriz de trazabilidad de endpoints JS hacia sus dependencias internas y su impacto potencial en BD "
        "(si el destino es SQL/HANA y las tablas son detectables desde el código)."
    )

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Endpoint origen (JS)"
    hdr[1].text = "Endpoint destino"
    hdr[2].text = "Handler destino"
    hdr[3].text = "Tablas/Vistas detectadas"

    for tr in traceability:
        row = table.add_row().cells
        row[0].text = tr["source"]
        row[1].text = tr["target"]
        row[2].text = tr["target_handler"]
        row[3].text = tr["db_tables"]

    doc.add_page_break()
    doc.add_heading("Resumen Ejecutivo Global", level=1)

    by_cat = {k: len(grouped.get(k, [])) for k in CATEGORY_ORDER}
    handler_count = defaultdict(int)
    for r in rows:
        handler_count[r["handler"]] += 1

    db_endpoints = sum(1 for r in rows if r["all_tables"])
    unique_tables = set()
    for r in rows:
        for t in r["all_tables"]:
            unique_tables.add(t)

    exec_text = (
        f"El análisis cubrió {len(rows)} endpoints del app portalclientescorporativos. "
        f"Distribución por categoría: "
        + ", ".join(f"{k}: {v}" for k, v in by_cat.items())
        + ". "
        + "Distribución por handler: "
        + ", ".join(f"{k}: {v}" for k, v in sorted(handler_count.items()))
        + ". "
        + f"Se identificó impacto directo o indirecto en BD en {db_endpoints} endpoint(s), "
        + f"con {len(unique_tables)} tabla(s)/vista(s) detectadas. "
        + "Los flujos con mayor dependencia transversal son: resumen corporativo, colaboradores y usuarios externos, "
        + "por su combinación de lógica JS, SQL/HANA y servicios SOAP heredados."
    )
    doc.add_paragraph(exec_text)

    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    app, eps = load_endpoints()
    grouped, rows, traceability = build_dataset(app, eps)
    write_doc(app, grouped, rows, traceability)
    print(f"Word generado en: {OUT_DOCX}")
